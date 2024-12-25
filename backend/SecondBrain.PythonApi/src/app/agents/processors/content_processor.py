from typing import Dict, Any, List, Optional, BinaryIO
import logging
from pathlib import Path
import fitz  # PyMuPDF
import docx
import pandas as pd
from bs4 import BeautifulSoup
import json
from ..core.agent_exceptions import ProcessingError
from ..utils.logging_utils import log_api_call

logger = logging.getLogger(__name__)

class ContentProcessor:
    """Content processor for different file types"""
    
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'process_pdf',
        '.docx': 'process_docx',
        '.xlsx': 'process_excel',
        '.csv': 'process_csv',
        '.json': 'process_json',
        '.html': 'process_html',
        '.txt': 'process_text'
    }
    
    @log_api_call
    def process_file(
        self,
        file_path: str,
        file_type: Optional[str] = None,
        **kwargs
    ) -> Dict[str, Any]:
        """Process file based on type"""
        try:
            path = Path(file_path)
            if not path.exists():
                raise ProcessingError(f"File not found: {file_path}")
                
            # Determine file type
            extension = path.suffix.lower()
            if file_type:
                extension = f".{file_type.lower()}"
                
            # Get processing method
            processor = getattr(
                self,
                self.SUPPORTED_EXTENSIONS.get(extension, 'process_text'),
                self.process_text
            )
            
            # Process file
            return processor(file_path, **kwargs)
            
        except Exception as e:
            error_msg = f"Failed to process file {file_path}: {str(e)}"
            logger.error(error_msg)
            raise ProcessingError(error_msg)
            
    def process_pdf(
        self,
        file_path: str,
        start_page: int = 0,
        end_page: Optional[int] = None,
        extract_images: bool = False
    ) -> Dict[str, Any]:
        """Process PDF file"""
        try:
            doc = fitz.open(file_path)
            content = {
                "text": "",
                "pages": [],
                "metadata": doc.metadata,
                "total_pages": len(doc)
            }
            
            # Process pages
            for page_num in range(start_page, end_page or len(doc)):
                page = doc[page_num]
                page_content = {
                    "number": page_num + 1,
                    "text": page.get_text(),
                    "images": []
                }
                
                # Extract images if requested
                if extract_images:
                    for img_index, img in enumerate(page.get_images()):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        if base_image:
                            page_content["images"].append({
                                "index": img_index,
                                "type": base_image["ext"],
                                "size": len(base_image["image"]),
                                "width": base_image.get("width"),
                                "height": base_image.get("height")
                            })
                            
                content["pages"].append(page_content)
                content["text"] += page_content["text"] + "\n\n"
                
            return content
            
        except Exception as e:
            raise ProcessingError(f"PDF processing failed: {str(e)}")
            
    def process_docx(
        self,
        file_path: str,
        include_headers: bool = True,
        include_footers: bool = True
    ) -> Dict[str, Any]:
        """Process DOCX file"""
        try:
            doc = docx.Document(file_path)
            content = {
                "text": "",
                "paragraphs": [],
                "tables": [],
                "headers": [],
                "footers": []
            }
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content["paragraphs"].append({
                        "text": para.text,
                        "style": para.style.name
                    })
                    content["text"] += para.text + "\n\n"
                    
            # Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                content["tables"].append(table_data)
                
            # Process headers if requested
            if include_headers:
                for section in doc.sections:
                    if section.header:
                        content["headers"].append(
                            section.header.paragraphs[0].text
                        )
                        
            # Process footers if requested
            if include_footers:
                for section in doc.sections:
                    if section.footer:
                        content["footers"].append(
                            section.footer.paragraphs[0].text
                        )
                        
            return content
            
        except Exception as e:
            raise ProcessingError(f"DOCX processing failed: {str(e)}")
            
    def process_excel(
        self,
        file_path: str,
        sheet_name: Optional[str] = None,
        max_rows: Optional[int] = None
    ) -> Dict[str, Any]:
        """Process Excel file"""
        try:
            content = {"sheets": {}}
            
            # Read Excel file
            if sheet_name:
                df = pd.read_excel(file_path, sheet_name=sheet_name, nrows=max_rows)
                content["sheets"][sheet_name] = self._process_dataframe(df)
            else:
                all_sheets = pd.read_excel(file_path, sheet_name=None, nrows=max_rows)
                for name, df in all_sheets.items():
                    content["sheets"][name] = self._process_dataframe(df)
                    
            return content
            
        except Exception as e:
            raise ProcessingError(f"Excel processing failed: {str(e)}")
            
    def process_csv(
        self,
        file_path: str,
        delimiter: str = ',',
        max_rows: Optional[int] = None
    ) -> Dict[str, Any]:
        """Process CSV file"""
        try:
            df = pd.read_csv(file_path, delimiter=delimiter, nrows=max_rows)
            return self._process_dataframe(df)
            
        except Exception as e:
            raise ProcessingError(f"CSV processing failed: {str(e)}")
            
    def process_json(
        self,
        file_path: str,
        encoding: str = 'utf-8'
    ) -> Dict[str, Any]:
        """Process JSON file"""
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return json.load(f)
                
        except Exception as e:
            raise ProcessingError(f"JSON processing failed: {str(e)}")
            
    def process_html(
        self,
        file_path: str,
        encoding: str = 'utf-8'
    ) -> Dict[str, Any]:
        """Process HTML file"""
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                soup = BeautifulSoup(f, 'html.parser')
                
            content = {
                "title": soup.title.string if soup.title else None,
                "text": soup.get_text(),
                "links": [
                    {"text": a.text, "href": a.get('href')}
                    for a in soup.find_all('a')
                ],
                "images": [
                    {"src": img.get('src'), "alt": img.get('alt')}
                    for img in soup.find_all('img')
                ],
                "metadata": {
                    meta.get('name', meta.get('property')): meta.get('content')
                    for meta in soup.find_all('meta')
                    if meta.get('name') or meta.get('property')
                }
            }
            
            return content
            
        except Exception as e:
            raise ProcessingError(f"HTML processing failed: {str(e)}")
            
    def process_text(
        self,
        file_path: str,
        encoding: str = 'utf-8'
    ) -> Dict[str, Any]:
        """Process text file"""
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
                
            return {
                "text": text,
                "lines": text.splitlines(),
                "size": len(text)
            }
            
        except Exception as e:
            raise ProcessingError(f"Text processing failed: {str(e)}")
            
    def _process_dataframe(self, df: pd.DataFrame) -> Dict[str, Any]:
        """Process pandas DataFrame into standard format"""
        return {
            "columns": list(df.columns),
            "data": df.values.tolist(),
            "shape": df.shape,
            "dtypes": df.dtypes.astype(str).to_dict()
        } 